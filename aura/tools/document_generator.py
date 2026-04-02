"""Document Generator Tool — create Word documents and PDFs from text or markdown.

Converts AURA's text output into professional formatted documents.
Saves to Desktop by default or a specified path.

Supported formats:
- Word (.docx) via python-docx
- PDF via fpdf2
- Markdown conversion to both formats
"""

import logging
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Dict, Any

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from fpdf import FPDF
    FPDF_AVAILABLE = True
except ImportError:
    FPDF_AVAILABLE = False

DESKTOP = Path.home() / "Desktop"


def _default_output(filename: str) -> str:
    """Generate a timestamped output path on Desktop."""
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    return str(DESKTOP / f"{ts}_{filename}")


class DocumentGeneratorTool:
    """Generate Word documents and PDFs from text or markdown — meeting notes, reports, proposals."""

    name = "document_generator"
    description = "Generate Word (.docx) and PDF documents from text or markdown — reports, meeting notes, proposals"

    # ------------------------------------------------------------------ #
    # Markdown parsing helpers
    # ------------------------------------------------------------------ #

    def _parse_markdown_blocks(self, text: str) -> List[Dict]:
        """Parse markdown text into semantic blocks."""
        blocks = []
        lines = text.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i]

            # Heading
            h_match = re.match(r"^(#{1,6})\s+(.+)", line)
            if h_match:
                level = len(h_match.group(1))
                blocks.append({"type": "heading", "level": level, "text": h_match.group(2).strip()})
                i += 1
                continue

            # Horizontal rule
            if re.match(r"^[-*_]{3,}\s*$", line):
                blocks.append({"type": "hr"})
                i += 1
                continue

            # Unordered list item
            li_match = re.match(r"^[-*+]\s+(.+)", line)
            if li_match:
                items = []
                while i < len(lines) and re.match(r"^[-*+]\s+(.+)", lines[i]):
                    items.append(re.match(r"^[-*+]\s+(.+)", lines[i]).group(1))
                    i += 1
                blocks.append({"type": "list", "ordered": False, "items": items})
                continue

            # Ordered list item
            oli_match = re.match(r"^\d+\.\s+(.+)", line)
            if oli_match:
                items = []
                while i < len(lines) and re.match(r"^\d+\.\s+(.+)", lines[i]):
                    items.append(re.match(r"^\d+\.\s+(.+)", lines[i]).group(1))
                    i += 1
                blocks.append({"type": "list", "ordered": True, "items": items})
                continue

            # Code block
            if line.strip().startswith("```"):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                i += 1  # skip closing ```
                blocks.append({"type": "code", "text": "\n".join(code_lines)})
                continue

            # Table
            if "|" in line and i + 1 < len(lines) and re.match(r"^\|[-| :]+\|$", lines[i + 1].strip()):
                table_lines = []
                while i < len(lines) and "|" in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                blocks.append({"type": "table", "raw": table_lines})
                continue

            # Empty line
            if not line.strip():
                i += 1
                continue

            # Paragraph
            para_lines = []
            while i < len(lines) and lines[i].strip() and not re.match(r"^[#\-*+|]|^\d+\.", lines[i]):
                para_lines.append(lines[i])
                i += 1
            if para_lines:
                blocks.append({"type": "paragraph", "text": " ".join(para_lines)})

        return blocks

    def _strip_inline_md(self, text: str) -> str:
        """Remove inline markdown (bold, italic, code, links) for plain text."""
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        text = re.sub(r"`(.+?)`", r"\1", text)
        text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
        return text

    # ------------------------------------------------------------------ #
    # Word document
    # ------------------------------------------------------------------ #

    def create_word_doc(
        self,
        title: str,
        content: str,
        output_path: Optional[str] = None,
        author: Optional[str] = None,
        subtitle: Optional[str] = None,
    ) -> Dict:
        """Create a formatted Word document.

        Args:
            title: Document title
            content: Body text (plain text or markdown)
            output_path: Where to save. Defaults to Desktop.
            author: Author name for metadata
            subtitle: Optional subtitle below title
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed. Run: pip install python-docx"}

        out = Path(output_path) if output_path else Path(_default_output(f"{title[:30].replace(' ', '_')}.docx"))
        out.parent.mkdir(parents=True, exist_ok=True)

        try:
            doc = Document()

            # Document properties
            props = doc.core_properties
            props.title = title
            if author:
                props.author = author
            props.created = datetime.now()

            # Page margins
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.2)
                section.right_margin = Inches(1.2)

            # Title
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if subtitle:
                sub = doc.add_paragraph(subtitle)
                sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sub.runs[0].italic = True

            # Date line
            date_para = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            doc.add_paragraph()  # spacer

            # Parse and render content
            blocks = self._parse_markdown_blocks(content)
            for block in blocks:
                btype = block["type"]

                if btype == "heading":
                    level = min(block["level"], 4)
                    doc.add_heading(self._strip_inline_md(block["text"]), level=level)

                elif btype == "paragraph":
                    p = doc.add_paragraph(self._strip_inline_md(block["text"]))
                    p.paragraph_format.space_after = Pt(6)

                elif btype == "list":
                    for item in block["items"]:
                        style = "List Number" if block["ordered"] else "List Bullet"
                        doc.add_paragraph(self._strip_inline_md(item), style=style)

                elif btype == "code":
                    p = doc.add_paragraph(block["text"])
                    p.style = "No Spacing"
                    run = p.runs[0] if p.runs else p.add_run(block["text"])
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)

                elif btype == "hr":
                    doc.add_paragraph("─" * 60)

                elif btype == "table":
                    self._add_markdown_table(doc, block["raw"])

            # Insert Table of Contents if document has 3+ headings
            self._maybe_insert_toc(doc, blocks)

            doc.save(str(out))
            return {
                "success": True,
                "format": "docx",
                "path": str(out),
                "title": title,
                "size_kb": round(out.stat().st_size / 1024, 1),
            }
        except Exception as e:
            logger.error(f"[DocGen] Word creation failed: {e}")
            return {"success": False, "error": str(e)}

    def _maybe_insert_toc(self, doc, blocks: List[Dict]):
        """Insert a Table of Contents after the title if there are 3+ headings."""
        headings = [b for b in blocks if b["type"] == "heading"]
        if len(headings) < 3:
            return
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            # Find insertion point: after the first few paragraphs (title, subtitle, date, spacer)
            # We insert before the first content paragraph
            insert_idx = min(4, len(doc.paragraphs))

            # Add TOC heading
            toc_heading = doc.add_paragraph('Table of Contents', style='TOC Heading')

            # Add a TOC field code (Word will render it when opened)
            p = doc.add_paragraph()
            run = p.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar1)

            run2 = p.add_run()
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
            run2._r.append(instrText)

            run3 = p.add_run()
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            run3._r.append(fldChar2)

            # Placeholder text for the TOC entries (replaced when Word updates fields)
            for h in headings:
                indent = "  " * (h["level"] - 1)
                run_entry = p.add_run(f'{indent}{h["text"]}\n')
                run_entry.font.size = Pt(10)

            run4 = p.add_run()
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')
            run4._r.append(fldChar3)

            doc.add_paragraph()  # spacer after TOC

            # Move TOC heading and TOC paragraph to the insertion point
            body = doc.element.body
            toc_elements = [toc_heading._p, p._p, doc.paragraphs[-1]._p]
            ref_element = doc.paragraphs[insert_idx]._p if insert_idx < len(doc.paragraphs) else None
            for elem in toc_elements:
                body.remove(elem)
                if ref_element is not None:
                    ref_element.addprevious(elem)
                else:
                    body.append(elem)
        except Exception as e:
            logger.debug(f"[DocGen] TOC insertion failed: {e}")

    def _add_markdown_table(self, doc, raw_lines: List[str]):
        """Add a markdown table to a Word doc."""
        try:
            rows = []
            for line in raw_lines:
                if re.match(r"^\|[-| :]+\|$", line.strip()):
                    continue  # skip separator
                cells = [c.strip() for c in line.strip().strip("|").split("|")]
                if cells:
                    rows.append(cells)
            if not rows:
                return
            max_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=max_cols)
            table.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                for c_idx, cell_text in enumerate(row):
                    if c_idx < max_cols:
                        cell = table.cell(r_idx, c_idx)
                        cell.text = self._strip_inline_md(cell_text)
                        if r_idx == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
        except Exception as e:
            logger.debug(f"[DocGen] Table render failed: {e}")

    # ------------------------------------------------------------------ #
    # PDF
    # ------------------------------------------------------------------ #

    def create_pdf(
        self,
        title: str,
        content: str,
        output_path: Optional[str] = None,
        author: Optional[str] = None,
    ) -> Dict:
        """Create a PDF document.

        Args:
            title: Document title
            content: Body text (markdown rendered as plain text)
            output_path: Where to save. Defaults to Desktop.
            author: Author name
        """
        if not FPDF_AVAILABLE:
            return {"success": False, "error": "fpdf2 not installed. Run: pip install fpdf2"}

        out = Path(output_path) if output_path else Path(_default_output(f"{title[:30].replace(' ', '_')}.pdf"))
        out.parent.mkdir(parents=True, exist_ok=True)

        try:
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=20)
            pdf.add_page()

            # Title
            pdf.set_font("Helvetica", "B", 20)
            pdf.set_text_color(30, 30, 30)
            pdf.cell(0, 12, title, ln=True, align="C")
            pdf.ln(2)

            # Date
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(130, 130, 130)
            pdf.cell(0, 8, datetime.now().strftime("%B %d, %Y"), ln=True, align="C")
            if author:
                pdf.cell(0, 8, f"By {author}", ln=True, align="C")
            pdf.ln(8)
            pdf.set_draw_color(200, 200, 200)
            pdf.line(20, pdf.get_y(), 190, pdf.get_y())
            pdf.ln(6)

            # Content
            blocks = self._parse_markdown_blocks(content)
            for block in blocks:
                btype = block["type"]
                pdf.set_text_color(30, 30, 30)

                if btype == "heading":
                    level = block["level"]
                    sizes = {1: 16, 2: 14, 3: 12, 4: 11, 5: 10, 6: 10}
                    pdf.set_font("Helvetica", "B", sizes.get(level, 12))
                    pdf.ln(4)
                    pdf.multi_cell(0, 8, self._strip_inline_md(block["text"]))
                    pdf.ln(2)

                elif btype == "paragraph":
                    pdf.set_font("Helvetica", "", 11)
                    pdf.multi_cell(0, 6, self._strip_inline_md(block["text"]))
                    pdf.ln(3)

                elif btype == "list":
                    pdf.set_font("Helvetica", "", 11)
                    for idx, item in enumerate(block["items"]):
                        bullet = f"{idx + 1}." if block["ordered"] else "•"
                        pdf.set_x(25)
                        pdf.multi_cell(0, 6, f"{bullet}  {self._strip_inline_md(item)}")
                    pdf.ln(2)

                elif btype == "code":
                    pdf.set_font("Courier", "", 9)
                    pdf.set_fill_color(245, 245, 245)
                    pdf.set_x(20)
                    lines = block["text"].splitlines()
                    for line in lines[:50]:  # cap at 50 lines
                        safe_line = line.encode("latin-1", errors="replace").decode("latin-1")
                        pdf.cell(0, 5, safe_line, ln=True, fill=True)
                    pdf.ln(3)

                elif btype == "hr":
                    pdf.set_draw_color(200, 200, 200)
                    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
                    pdf.ln(4)

            pdf.output(str(out))
            return {
                "success": True,
                "format": "pdf",
                "path": str(out),
                "title": title,
                "size_kb": round(out.stat().st_size / 1024, 1),
            }
        except Exception as e:
            logger.error(f"[DocGen] PDF creation failed: {e}")
            return {"success": False, "error": str(e)}

    # ------------------------------------------------------------------ #
    # Simple docx from text
    # ------------------------------------------------------------------ #

    def create_docx(self, text: str, output_path: str) -> Dict:
        """Create a simple Word document from text with basic markdown formatting.

        Args:
            text: Body text (supports # headings and - bullet points)
            output_path: Where to save the .docx file
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed. Run: pip install python-docx"}

        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)

        try:
            doc = Document()
            doc.add_heading("AURA Document", level=0)

            for line in text.splitlines():
                stripped = line.strip()
                if not stripped:
                    continue

                # Headings
                h_match = re.match(r"^(#{1,6})\s+(.+)", stripped)
                if h_match:
                    level = min(len(h_match.group(1)), 4)
                    doc.add_heading(h_match.group(2).strip(), level=level)
                    continue

                # Bullet points
                li_match = re.match(r"^[-*+]\s+(.+)", stripped)
                if li_match:
                    doc.add_paragraph(li_match.group(1), style="List Bullet")
                    continue

                # Regular paragraph
                doc.add_paragraph(stripped)

            doc.save(str(out))
            return {
                "success": True,
                "format": "docx",
                "path": str(out),
                "size_kb": round(out.stat().st_size / 1024, 1),
            }
        except Exception as e:
            logger.error(f"[DocGen] create_docx failed: {e}")
            return {"success": False, "error": str(e)}

    # ------------------------------------------------------------------ #
    # Convenience wrappers
    # ------------------------------------------------------------------ #

    def markdown_to_word(self, markdown_text: str, output_path: Optional[str] = None, title: Optional[str] = None) -> Dict:
        """Convert markdown directly to a Word document."""
        t = title or "Document"
        first_h1 = re.search(r"^#\s+(.+)", markdown_text, re.MULTILINE)
        if first_h1 and not title:
            t = first_h1.group(1)
            markdown_text = markdown_text[first_h1.end():].strip()
        return self.create_word_doc(t, markdown_text, output_path)

    def markdown_to_pdf(self, markdown_text: str, output_path: Optional[str] = None, title: Optional[str] = None) -> Dict:
        """Convert markdown directly to a PDF."""
        t = title or "Document"
        first_h1 = re.search(r"^#\s+(.+)", markdown_text, re.MULTILINE)
        if first_h1 and not title:
            t = first_h1.group(1)
            markdown_text = markdown_text[first_h1.end():].strip()
        return self.create_pdf(t, markdown_text, output_path)

    def create_meeting_notes(self, content: str, meeting_name: Optional[str] = None, format: str = "docx") -> Dict:
        """Create a formatted meeting notes document."""
        title = meeting_name or f"Meeting Notes — {datetime.now().strftime('%B %d, %Y')}"
        if format.lower() == "pdf":
            return self.create_pdf(title, content)
        return self.create_word_doc(title, content)

    def execute(self, action: str, **kwargs) -> Dict:
        """Execute a document generation action."""
        a = action.lower().strip()
        title = kwargs.get("title") or "Document"
        content = kwargs.get("content") or kwargs.get("text") or ""
        output = kwargs.get("output_path") or kwargs.get("path")
        fmt = kwargs.get("format", "docx").lower()

        if "meeting" in a or "notes" in a:
            return self.create_meeting_notes(content, title, fmt)
        if "markdown" in a or "md" in a:
            if fmt == "pdf":
                return self.markdown_to_pdf(content, output, title)
            return self.markdown_to_word(content, output, title)
        if fmt == "pdf" or "pdf" in a:
            return self.create_pdf(title, content, output, kwargs.get("author"))
        return self.create_word_doc(title, content, output, kwargs.get("author"), kwargs.get("subtitle"))
